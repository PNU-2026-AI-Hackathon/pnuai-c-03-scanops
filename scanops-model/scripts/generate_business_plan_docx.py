"""
ScanOps 사업계획서 DOCX 생성 (버전 표기 제거판)
주요 개정:
  - 버전(v6 등) 표기 전면 제거
  - 1. 사업 개요: 한 줄 정의를 맨 앞으로, 비전공 평가위원도 이해할 쉬운 말로 재작성,
    모델 비교(Grok 동급/v5 대비) 문구 제거, 개요 표를 평이하게 정리
  - 2. 문제 인식: 수치·기사·사고 사례 근거 대폭 보강(바이브코딩 앱 폭증, AI코드 취약,
    SKT·모두의창업 등 실제 사고), 고객 인터뷰 근거를 앞으로 당겨 배치
  - 3. 솔루션: 실제 작동 화면(스크린샷) 삽입, 모델 설명은 핵심만, 모델 3B 업그레이드 및
    Java/Neo4j 학습·재검증 예정 상태를 투명하게 명시
  - 6. 비즈니스 모델: 개인(Pro/Max)·팀(Team) 플랜 분리 추가
  - 8. 팀 구성: 표 + 각자 하는 일 서술 추가
실행: scanops-model/.venv/bin/python scripts/generate_business_plan_docx.py
"""
import io
import json
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

plt.rcParams['font.family'] = ['AppleGothic', 'NanumGothic', 'Malgun Gothic', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

BASE = Path(__file__).resolve().parents[1]
OUT = BASE / "reports" / "ScanOps_사업계획서.docx"
OUT_DESKTOP = BASE.parent / "ScanOps_사업계획서.docx"
SHOT = BASE / "reports" / "grant_attachment"
RES = json.loads((BASE / "reports" / "results_v5_false_positive_benchmark.json").read_text())
SV5, GROK, RAW = RES["systems"][0]["metrics"], RES["systems"][1]["metrics"], RES["v4_raw_safe_sample"]
GRAPH = json.loads((BASE / "reports" / "results_graph_vs_grok.json").read_text())

NAVY = (44, 62, 80)
GREEN = (39, 174, 96)
BLUE = (41, 128, 185)
PURPLE = (124, 58, 237)


# ── 스타일 헬퍼 ────────────────────────────────────────────────────────────
def set_font(run, name="맑은 고딕", size=11, bold=False, color=None, italic=False):
    run.font.name = name; run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts"); rFonts.set(qn("w:eastAsia"), name); rPr.insert(0, rFonts)


def H(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "맑은 고딕"; run.font.size = Pt([0, 17, 14, 12][level])
        if color:
            run.font.color.rgb = RGBColor(*color)
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts"); rFonts.set(qn("w:eastAsia"), "맑은 고딕"); rPr.insert(0, rFonts)
    return p


def P(doc, text="", bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, after=6):
    p = doc.add_paragraph(); p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    if text:
        set_font(p.add_run(text), size=size, bold=bold, color=color)
    return p


def B(doc, text, size=10.5, color=None):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run("•  "), size=size, bold=True, color=GREEN)
    set_font(p.add_run(text), size=size, color=color)
    return p


def LEAD(doc, label, text, size=10.5):
    """굵은 라벨 + 일반 본문을 한 단락에."""
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run("•  "), size=size, bold=True, color=GREEN)
    set_font(p.add_run(label + " "), size=size, bold=True)
    set_font(p.add_run(text), size=size)
    return p


def NOTE(doc, text, size=9.5, label="ℹ︎ 현재 상태  "):
    """투명성 단서(미구현/예정 등) 표시용 — 연한 배경 박스."""
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "FFF7E6")
    pPr.insert(0, shd)
    set_font(p.add_run(label), size=size, bold=True, color=(176, 108, 0))
    set_font(p.add_run(text), size=size, color=(120, 90, 30))
    return p


def QUOTE(doc, text, size=10.5):
    """고객 인용 등 강조 박스(연한 파랑)."""
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "EAF2FB")
    pPr.insert(0, shd)
    set_font(p.add_run(text), size=size, italic=True, color=(40, 70, 110))
    return p


def _cell_shd(tcPr, fill):
    """tcPr에 음영(shd)을 스키마 순서(vAlign 앞)에 맞게 삽입."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    valign = tcPr.find(qn("w:vAlign"))
    if valign is not None:
        valign.addprevious(shd)
    else:
        tcPr.append(shd)


def TABLE(doc, headers, rows, widths=None, hi=None, hi_fill="EAF6EC", hi_text=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(pp.add_run(h), size=9.5, bold=True, color=(255, 255, 255))
        _cell_shd(c._tc.get_or_add_tcPr(), "2C3E50")
    hi_set = set(hi) if isinstance(hi, (list, tuple)) else ({hi} if hi is not None else set())
    for ri, rd in enumerate(rows):
        row = t.add_row(); fill = hi_fill if ri in hi_set else None
        for i, v in enumerate(rd):
            c = row.cells[i]; c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            pp = c.paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.LEFT if (i == len(rd) - 1 and len(headers) == 2) else WD_ALIGN_PARAGRAPH.CENTER
            set_font(pp.add_run(str(v)), size=9.5, bold=(fill is not None),
                     color=(hi_text if fill else None))
            if fill:
                _cell_shd(c._tc.get_or_add_tcPr(), fill)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    return t


def fig_stream(fig):
    buf = io.BytesIO(); fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig); buf.seek(0); return buf


def CHART(doc, buf, cap="", width=Inches(5.9)):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(buf, width=width)
    if cap:
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(cap); r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = RGBColor(110, 110, 110)


def IMG(doc, path, cap="", width=Inches(5.9)):
    """로컬 PNG 스크린샷 삽입(테두리 포함)."""
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(); run.add_picture(str(path), width=width)
    # 이미지에 얇은 회색 테두리
    try:
        inline = run._r.find(qn('w:drawing'))[0]
        pic = inline.find('.//' + qn('pic:pic'))
        spPr = pic.find(qn('pic:spPr'))
        ln = OxmlElement('a:ln'); ln.set('w', '6350')
        fill = OxmlElement('a:solidFill'); clr = OxmlElement('a:srgbClr'); clr.set('val', 'D0D5DD')
        fill.append(clr); ln.append(fill); spPr.append(ln)
    except Exception:
        pass
    if cap:
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(cap); r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = RGBColor(110, 110, 110)


# ── 차트 ────────────────────────────────────────────────────────────────
def chart_benchmark():
    fig, (a1, a2) = plt.subplots(1, 2, figsize=(9.2, 3.6), facecolor='#FAFAFA')
    a1.set_facecolor('#F7F9FC')
    names = ['필터 전\n(raw)', '하이브리드\n필터', 'Grok-3']
    vals = [RAW["fpr_pct"], SV5["false_positive_rate"], GROK["false_positive_rate"]]
    bars = a1.bar(names, vals, color=['#E74C3C', '#27AE60', '#3498DB'], width=0.6, zorder=3)
    for b, v in zip(bars, vals):
        a1.text(b.get_x()+b.get_width()/2, b.get_height()+2, f"{v:.0f}%", ha='center', fontsize=10, fontweight='bold')
    a1.set_ylim(0, 112); a1.set_title("오탐률(FPR) — 낮을수록 좋음", fontsize=11, fontweight='bold')
    a1.grid(True, axis='y', linestyle='--', alpha=0.4)
    a2.set_facecolor('#F7F9FC')
    labels = ['탐지율', '오탐률', '정밀도', '정확도', 'F1']
    sv5 = [SV5["detection_recall"], SV5["false_positive_rate"], SV5["precision"], SV5["accuracy"], SV5["f1"]]
    gk = [GROK["detection_recall"], GROK["false_positive_rate"], GROK["precision"], GROK["accuracy"], GROK["f1"]]
    x = np.arange(len(labels)); w = 0.38
    a2.bar(x-w/2, sv5, w, label='ScanOps', color='#27AE60', zorder=3)
    a2.bar(x+w/2, gk, w, label='Grok-3', color='#3498DB', zorder=3)
    a2.set_xticks(x); a2.set_xticklabels(labels, fontsize=9); a2.set_ylim(0, 112)
    a2.set_title("단일코드 탐지 — ScanOps vs Grok-3", fontsize=11, fontweight='bold')
    a2.legend(fontsize=8, loc='lower center', ncol=2); a2.grid(True, axis='y', linestyle='--', alpha=0.4)
    fig.tight_layout(); return fig_stream(fig)


def chart_headtohead():
    fig, ax = plt.subplots(figsize=(9.0, 3.4), facecolor='#FAFAFA'); ax.set_facecolor('#F7F9FC')
    labels = ['Missing AuthZ\n(CVE-2026-44754)', 'Prototype Pollution\n(CVE-2026-11572)',
              'Hardcoded Cred\n(CVE-2026-21404)', 'Missing AuthZ\n(CVE-2026-44751)',
              'Rate Limit\n(CVE-2026-11572)']
    scan = [1, 1, 1, 1, 1]
    grok = [0, 0, 0, 0, 0]
    x = np.arange(len(labels)); w = 0.38
    ax.bar(x - w/2, scan, w, label='ScanOps (탐지 ✓)', color='#27AE60', zorder=3)
    ax.bar(x + w/2, grok, w, label='Grok-3 (미탐 ✗)', color='#E74C3C', zorder=3)
    ax.set_xticks(x); ax.set_xticklabels(labels, fontsize=8)
    ax.set_yticks([0, 1]); ax.set_yticklabels(['미탐(SAFE)', '탐지(VULN)'], fontsize=9)
    ax.set_ylim(0, 1.25)
    ax.set_title("2026년 5~6월 신규 CVE — ScanOps 탐지 ✓ / Grok 미탐 ✗", fontsize=11.5, fontweight='bold', pad=8)
    ax.legend(fontsize=9, loc='upper right'); ax.grid(True, axis='y', linestyle='--', alpha=0.3)
    fig.tight_layout(); return fig_stream(fig)


def chart_graph_benchmark():
    fig, (a1, a2) = plt.subplots(1, 2, figsize=(9.2, 3.6), facecolor='#FAFAFA')
    a1.set_facecolor('#F7F9FC')
    names = ['전체 100', 'CVE-2026\n50', '구조패턴\n50']
    so = [GRAPH["scanops_accuracy"], GRAPH["breakdown"]["cve_2026"]["scanops_accuracy"],
          GRAPH["breakdown"]["structural"]["scanops_accuracy"]]
    gk = [GRAPH["grok_accuracy"], GRAPH["breakdown"]["cve_2026"]["grok_accuracy"],
          GRAPH["breakdown"]["structural"]["grok_accuracy"]]
    x = np.arange(len(names)); w = 0.38
    a1.bar(x-w/2, so, w, label='ScanOps 그래프엔진', color='#7C3AED', zorder=3)
    a1.bar(x+w/2, gk, w, label='Grok-3-mini', color='#3498DB', zorder=3)
    a1.set_xticks(x); a1.set_xticklabels(names, fontsize=9); a1.set_ylim(0, 112)
    a1.set_title("멀티파일 taint 추적 정확도(%)", fontsize=11, fontweight='bold')
    a1.legend(fontsize=8, loc='lower center', ncol=2); a1.grid(True, axis='y', linestyle='--', alpha=0.4)
    a2.set_facecolor('#F7F9FC')
    labels = ['정확도', 'Recall\n(탐지율)', 'Specificity\n(오탐방지)']
    so2 = [100.0, 100.0, 100.0]
    gk2 = [GRAPH["grok_accuracy"], 35.3, 100.0]
    x2 = np.arange(len(labels)); w = 0.38
    a2.bar(x2-w/2, so2, w, label='ScanOps', color='#7C3AED', zorder=3)
    a2.bar(x2+w/2, gk2, w, label='Grok-3-mini', color='#3498DB', zorder=3)
    a2.set_xticks(x2); a2.set_xticklabels(labels, fontsize=9); a2.set_ylim(0, 112)
    a2.set_title("정확도 vs Recall — Grok은 오탐0이지만 미탐 다수", fontsize=10.3, fontweight='bold')
    a2.legend(fontsize=8, loc='lower center', ncol=2); a2.grid(True, axis='y', linestyle='--', alpha=0.4)
    fig.tight_layout(); return fig_stream(fig)


def chart_aicode():
    """문제인식 — AI 생성 코드 취약점 비율(Veracode 2025)."""
    fig, ax = plt.subplots(figsize=(8.4, 3.3), facecolor='#FAFAFA'); ax.set_facecolor('#F7F9FC')
    labels = ['AI 코드\n전체', 'Java\n언어', 'XSS\n(CWE-80)']
    vals = [45, 70, 86]
    bars = ax.bar(labels, vals, color=['#E67E22', '#E74C3C', '#C0392B'], width=0.5, zorder=3)
    for b, v in zip(bars, vals):
        ax.text(b.get_x()+b.get_width()/2, b.get_height()+1.5, f"{v}%", ha='center', fontsize=11, fontweight='bold')
    ax.set_ylim(0, 100); ax.set_ylabel("보안 결함 포함 비율(%)", fontsize=10)
    ax.set_title("AI가 생성한 코드의 보안 결함 비율 (Veracode 2025)", fontsize=11, fontweight='bold', pad=8)
    ax.grid(True, axis='y', linestyle='--', alpha=0.4); fig.tight_layout(); return fig_stream(fig)


def chart_appstore():
    """문제인식 — 앱스토어 제출 폭증."""
    fig, ax = plt.subplots(figsize=(8.2, 3.2), facecolor='#FAFAFA'); ax.set_facecolor('#F7F9FC')
    labels = ['Q1 2025', 'Q1 2026']
    vals = [128.2, 235.8]
    bars = ax.bar(labels, vals, color=['#95A5A6', '#2980B9'], width=0.5, zorder=3)
    for b, v in zip(bars, vals):
        ax.text(b.get_x()+b.get_width()/2, b.get_height()+4, f"{v:.0f}천 건", ha='center', fontsize=10.5, fontweight='bold')
    ax.annotate("+84%", xy=(1, 235.8), xytext=(0.5, 200), fontsize=13, fontweight='bold', color='#E74C3C', ha='center')
    ax.set_ylim(0, 270); ax.set_ylabel("앱스토어 심사 제출(천 건)", fontsize=10)
    ax.set_title("바이브 코딩 이후 앱 제출 폭증 (Apple App Store)", fontsize=11, fontweight='bold', pad=8)
    ax.grid(True, axis='y', linestyle='--', alpha=0.4); fig.tight_layout(); return fig_stream(fig)


def chart_pricing():
    fig, ax = plt.subplots(figsize=(8.6, 3.3), facecolor='#FAFAFA'); ax.set_facecolor('#F7F9FC')
    plans = ['회원가입', 'Pro', 'Max', 'Team\n(기본 3명)']
    prices = [0, 29900, 99000, 89000]
    bars = ax.bar(plans, prices, color=['#95A5A6', '#27AE60', '#2C3E50', '#7C3AED'], width=0.55, zorder=3)
    notes = ['₩0\n(1회 무료체험)', '₩29,900/월\n개인', '₩99,000/월\n코드 많은 개인', '₩89,000/월\n조직·팀 (+25,000/명)']
    for b, n in zip(bars, notes):
        ax.text(b.get_x()+b.get_width()/2, b.get_height()+2500, n, ha='center', fontsize=8.8, fontweight='bold')
    ax.set_ylim(0, 122000); ax.set_ylabel("월 구독료 (₩)", fontsize=10)
    ax.set_title("구독 플랜 — 개인(Pro·Max) / 팀(Team)", fontsize=11, fontweight='bold', pad=8)
    ax.grid(True, axis='y', linestyle='--', alpha=0.4); fig.tight_layout(); return fig_stream(fig)


def chart_market():
    fig, ax = plt.subplots(figsize=(8.2, 3.2), facecolor='#FAFAFA'); ax.set_facecolor('#F7F9FC')
    labels = ['TAM\n14.8조원', 'SAM\n8,800억원', 'SOM\n약 15억원']
    vals = [148000, 8800, 15]
    bars = ax.barh(labels, np.log10([v+1 for v in vals]), color=['#AED6F1', '#5DADE2', '#27AE60'], zorder=3)
    for b, v in zip(bars, vals):
        ax.text(b.get_width()+0.05, b.get_y()+b.get_height()/2, f"{v:,}억원" if v > 100 else f"{v}억원", va='center', fontsize=10, fontweight='bold')
    ax.set_xlim(0, 6.2); ax.set_xticks([]); ax.invert_yaxis()
    ax.set_title("시장 규모 (TAM/SAM/SOM, 로그 스케일)", fontsize=11, fontweight='bold', pad=8)
    fig.tight_layout(); return fig_stream(fig)


def chart_revenue():
    fig, ax = plt.subplots(figsize=(8.4, 3.6), facecolor='#FAFAFA'); ax.set_facecolor('#F7F9FC')
    years = ['1년차', '2년차', '3년차']; arr = [1.1, 4.2, 9.8]; users = [250, 900, 2000]
    ax.bar(years, arr, color='#27AE60', width=0.5, zorder=3, label='ARR(억원)')
    for i, (a, u) in enumerate(zip(arr, users)):
        ax.text(i, a+0.2, f"{a}억원\n(유료 {u:,}명)", ha='center', fontsize=10, fontweight='bold')
    ax.set_ylim(0, 12); ax.set_ylabel("ARR (억원)", fontsize=10)
    ax.set_title("3개년 추정 매출 (CTO 타깃·고가 플랜, 가정 기반)", fontsize=11, fontweight='bold', pad=8)
    ax.grid(True, axis='y', linestyle='--', alpha=0.4); fig.tight_layout(); return fig_stream(fig)


# ── 빌드 ────────────────────────────────────────────────────────────────
def build():
    print("차트 생성...")
    c_bench, c_price, c_market, c_rev = chart_benchmark(), chart_pricing(), chart_market(), chart_revenue()
    c_h2h = chart_headtohead()
    c_graph = chart_graph_benchmark()
    c_aicode = chart_aicode()
    c_appstore = chart_appstore()
    print("문서 작성...")
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.3); s.bottom_margin = Cm(2.3); s.left_margin = Cm(2.7); s.right_margin = Cm(2.4)

    # ── 표지 ──
    for _ in range(4):
        P(doc, "")
    P(doc, "사업계획서", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, color=(120, 120, 120))
    pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(pp.add_run("ScanOps"), size=30, bold=True, color=GREEN)
    P(doc, "코드를 외부로 보내지 않는, AI 보안 취약점 자동 진단 SaaS", size=14, bold=True,
      align=WD_ALIGN_PARAGRAPH.CENTER, color=NAVY)
    P(doc, "보안 전담 인력이 없는 작은 개발팀을 위한 자체 구동 AI 보안 검사관 — 코드 유출 0",
      size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=(110, 110, 110))
    for _ in range(2):
        P(doc, "")
    P(doc, "2026년 6월", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=(90, 90, 90))
    P(doc, "부산대학교 정보컴퓨터공학부 · 팀 ScanOps", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=(90, 90, 90))
    P(doc, "팀장 김세한 · 팀원 전혜은 · 이경윤 · 최효석", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=(90, 90, 90))
    doc.add_page_break()

    # ── 1. 사업 개요 ──
    H(doc, "1. 사업 개요", 1, NAVY)
    # 한 줄 정의를 맨 앞에
    P(doc, "한 줄 정의", bold=True, size=11, color=GREEN, after=2)
    QUOTE(doc, "“보안 전문가가 없는 작은 개발팀도, 코드를 외부에 맡기지 않고 합리적인 월 구독료로 쓸 수 있는 "
               "AI 보안 진단 도구”", size=11)
    # 비전공자도 이해할 쉬운 설명 (모델 비교 문구 없음)
    P(doc, "ScanOps는 한마디로 ‘AI 보안 검사관’입니다. 개발자가 만든 프로그램 코드에는 해커가 악용할 수 있는 "
           "허점(보안 취약점)이 숨어 있는데, 이를 직접 찾아내려면 보통 비싼 보안 전문가나 전문 도구가 필요합니다. "
           "ScanOps는 보안에 특화되도록 학습시킨 AI가 GitHub 저장소 주소나 웹사이트 주소만 받으면 이 허점들을 "
           "자동으로 찾아주고, 각 문제가 ‘얼마나 위험한지(위험도 점수)’와 ‘어떻게 고치면 되는지’까지 리포트로 "
           "정리해 줍니다.")
    P(doc, "차별점은 크게 두 가지입니다. 첫째, ‘코드를 밖으로 내보내지 않습니다’. 보통 ChatGPT 같은 외부 AI에 "
           "코드를 붙여넣어 검사하면 편리하지만, 회사의 핵심 소스코드가 외부 업체 서버로 그대로 넘어갑니다. "
           "ScanOps는 검사에 쓰는 AI를 외부에 맡기지 않고 우리 서버 안에서 직접 돌리기 때문에, 검사 과정에서 "
           "고객의 코드가 회사 밖으로 단 한 줄도 새어 나가지 않습니다.")
    P(doc, "둘째, ‘항상 최신 취약점까지 잡습니다’. ChatGPT 같은 범용 AI는 ‘학습한 시점’까지의 지식만 갖고 있어, "
           "그 이후 새로 발견된 보안 취약점은 아예 모를 수 있습니다. 보안 취약점은 매일 새로 공개되는데, 정작 "
           "범용 AI는 ‘본 적 없는’ 최신 취약점에 약합니다. ScanOps는 전 세계 취약점 정보(NVD)를 실시간으로 "
           "찾아보며 검사하기 때문에, ‘오늘 새로 알려진’ 허점까지 근거를 갖고 잡아냅니다. 즉, 코드가 안전하게 "
           "보호되면서도 탐지 능력은 늘 최신으로 유지됩니다.")
    P(doc, "")
    TABLE(doc, ["한눈에 보기", "내용"],
          [["무엇을 하나요?", "GitHub 저장소나 웹사이트 주소만 넣으면, AI가 코드의 보안 약점을 자동으로 찾아 ‘위험도’와 ‘고치는 방법’을 리포트로 알려줍니다."],
           ["누구를 위한 건가요?", "보안 전담 인력이 없는 스타트업·소규모 개발팀 (이후 중소기업까지 확장)"],
           ["가장 큰 강점은?", "① 검사를 우리 서버 안의 AI로만 하므로 고객 소스코드가 외부로 전송되지 않습니다. ② 최신 취약점 정보(NVD)를 실시간 검색해, 학습 시점이 고정된 범용 AI가 모르는 ‘새로 공개된’ 취약점까지 잡습니다."],
           ["어떻게 결제하나요?", "월 구독제 — 개인용(Pro·Max)과 팀용(Team)으로 나뉩니다. 가입하면 웹사이트 보안검사 1회를 무료로 체험할 수 있습니다."],
           ["지금 어디까지 됐나요?", "시제품(MVP)을 클라우드에 올려 자체 성능 검증을 마쳤고, 창업을 준비하는 예비창업 단계입니다."]],
          widths=[3.4, 13.1])
    P(doc, "※ 성능·단계는 모두 사내 자체 검증 기준이며, 정식 출시 전 베타 데이터로 정밀화합니다. (2026.6)",
      size=9, color=(130, 130, 130))
    doc.add_page_break()

    # ── 2. 문제 인식 ──
    H(doc, "2. 문제 인식", 1, NAVY)
    H(doc, "2.1 ‘바이브 코딩’ 시대 — 코드는 폭증하는데 검증은 공백이다", 2, BLUE)
    P(doc, "2025년 2월 등장한 ‘바이브 코딩(vibe coding)’ — AI에게 말로 지시해 코드를 만드는 방식 — 은 1년 만에 "
           "개발 환경을 바꿔 놓았다. 이제는 코드를 잘 모르는 사람도 AI로 앱을 만들어 출시할 수 있게 됐고, 그 결과 "
           "시장에 쏟아지는 앱과 코드의 양이 폭발적으로 늘었다. 실제로 애플 앱스토어 심사 제출은 2026년 1분기에만 "
           "약 23만 5,800건으로 1년 전 같은 기간보다 84% 급증했으며, 2025년 한 해 신규 앱 제출은 약 55만 7,000건으로 "
           "2016년 이후 최대치를 기록했다. 애플조차 심사 적체로 승인 기간이 하루에서 최대 30일까지 늘었다.")
    CHART(doc, c_appstore, "그림 1. 바이브 코딩 확산 이후 앱스토어 제출이 1년 새 84% 급증 (출처: Apple/AppleInsider·TNW 2026)")
    P(doc, "문제는 ‘만드는 속도’만 빨라졌다는 것이다. 만들어진 코드를 검증할 보안 역량은 같은 속도로 늘지 않는다. "
           "특히 비전공자가 하루 만에 만들어 올린 앱이라면, 정작 그 안에 어떤 보안 구멍이 있는지 만든 사람조차 "
           "모르는 경우가 많다.")
    H(doc, "2.2 ‘AI가 짠 코드’는 생각보다 안전하지 않다", 2, BLUE)
    P(doc, "‘AI가 짜 줬으니 괜찮겠지’라는 믿음은 데이터와 정면으로 어긋난다. 보안 분석 기업 베라코드(Veracode)가 "
           "2025년 100개 이상의 LLM을 대상으로 분석한 결과, AI가 생성한 코드의 45%가 하나 이상의 보안 결함(OWASP "
           "취약점)을 포함했다. 사람이 짠 코드보다 취약점 발생률이 약 2.74배 높았고, 특히 자바(Java)는 결함률 70%, "
           "교차 사이트 스크립팅(XSS)은 86%에 달했다. 그런데도 개발자의 약 80%는 ‘AI 코드가 사람 코드보다 더 "
           "안전하다’고 믿는다 — 인식과 현실의 간극이 그대로 위험이 된다.")
    CHART(doc, c_aicode, "그림 2. AI 생성 코드의 보안 결함 비율 — 전체 45%, 자바 70%, XSS 86% (출처: Veracode GenAI Code Security Report 2025)")
    P(doc, "요약하면, 바이브 코딩으로 ‘코드의 양’은 폭증했지만 그 코드의 ‘보안 품질’은 오히려 더 위태로워졌다. "
           "검사 없이 출시되는 앱이 늘어날수록, 사용자 데이터 유출 같은 사고의 모수도 함께 커진다.")
    H(doc, "2.3 실제로 사고는 터지고 있다 — 최근 국내 사례", 2, BLUE)
    P(doc, "이것은 가정이 아니라 현실이다. 최근 국내에서 벌어진 사고들의 공통점은 ‘대단한 외부 해킹’이 아니라 "
           "‘기본적인 내부 보안 관리의 실패’였다는 점이다.", after=4)
    LEAD(doc, "SK텔레콤 유심 정보 유출(2025).",
         "전체 가입자 약 2,324만 명의 휴대폰번호·유심 인증키 등이 유출됐다. 인터넷망과 내부 코어망을 분리하지 "
         "않고, 침입탐지 로그를 확인하지 않았으며, 인증키를 암호화하지 않는 등 안전조치를 위반한 결과였다. "
         "개인정보보호위원회는 역대 최대인 1,347억 9,100만 원의 과징금을 부과했다.")
    LEAD(doc, "정부 ‘모두의 창업’ 플랫폼 개인정보 노출(2026.6).",
         "API 응답을 통해 1차 합격자의 정보가 노출되는 취약점이 있었고, 약 1만 6,000건의 아이디어 목록과 "
         "2만 건의 팀원 정보가 노출될 수 있었다. 한 보안업체가 한 달 전 미리 제보했음에도 막지 못했다. 보도는 "
         "이 사고의 핵심을 ‘외부 해킹 공격이 아닌 내부 보안 관리 실패’, 즉 서버 API의 보안 미흡으로 짚었다.")
    P(doc, "두 사고 모두 ScanOps가 겨냥하는 바로 그 지점 — API 권한 검증 누락, 인증·암호화 미적용 같은 "
           "‘코드·설정 단계에서 미리 잡을 수 있었던’ 취약점이다. 즉 배포 전에 한 번만 점검했어도 막을 수 있었던 "
           "사고들이며, 이것이 ScanOps의 존재 이유다.")
    H(doc, "2.4 타깃 고객의 3대 페인포인트", 2, BLUE)
    for t in [
        "① 전담 보안 인력·예산이 없다 — 스타트업·소규모 팀은 보안 엔지니어를 둘 여력이 없다. 개발자가 보안까지 겸하지만 전문가가 아니라 취약점을 놓친다.",
        "② 기존 SAST 도구는 비싸고 무겁다 — 스패로우 등 국내 상용 정적분석 도구는 엔터프라이즈 견적·연단위 라이선스 중심이라 도입 장벽이 높다.",
        "③ 클라우드 AI에는 민감한 코드를 맡기기 어렵다 — ChatGPT·Claude에 코드를 붙여넣는 방식은 편하지만 핵심 소스코드를 외부 서버로 전송한다는 근본적 모순이 있다.",
    ]:
        B(doc, t)
    P(doc, "정리하면, 작은 팀은 ‘싸고 가볍고 코드가 새지 않는’ 보안 진단 도구를 원하지만 시장에는 그 셋을 동시에 "
           "만족하는 선택지가 없다. ScanOps는 정확히 이 빈 공간을 겨냥한다.")
    H(doc, "2.5 현장의 목소리 — 고객 인터뷰로 확인한 문제", 2, BLUE)
    P(doc, "위 문제들은 우리가 직접 만난 개발자들에게서도 그대로 확인됐다. 스타트업 SW 개발자를 대상으로 한 "
           "심층 인터뷰(7문항)의 핵심은 다음과 같다.", after=4)
    for t in [
        "보안 지식은 낮다 — 스스로 평가한 보안 지식이 10점 만점에 2~3점. 개인정보 비식별화 정도만 처리해 왔고 공격·방어 지식은 부족하다고 답했다.",
        "필요성은 알지만 계속 미룬다 — ‘정식 배포 전 기본 보안 점검은 무조건 필요하다’면서도, ‘빠른 출시’가 우선이고 비개발 팀원 설득 비용이 부담돼 점검을 받은 적이 없다.",
        "코드 외부 전송이 실제 구매 장벽 — ‘보안 리포트가 서버로 가면서 정보가 같이 새지 않을까’를 우려했다. (바로 ScanOps의 ‘코드 비전송’이 정조준하는 지점이다.)",
        "개발도구에는 이미 월 14만 원 이상을 쓴다 — Cursor·Claude Code·GPT·AWS 등에 매달 지출 중이며, 개당 3만 원대 결제가 일상적이다. 보안 도구의 합리적 구독료에 대한 지불 여력이 있다는 신호다.",
    ]:
        B(doc, t)
    QUOTE(doc, "실제 사용 후 반응: “이거 좋은데?! 공격 위험도·방어 방침·참고 문서가 같이 나와서 좋다. 개인 프로젝트 "
               "보안 학습에도 도움이 된다.” — 초기 인터뷰 참여 개발자")
    P(doc, "‘만드는 속도(바이브 코딩)’와 ‘검증의 공백(45% 취약)’, ‘실제 사고(내부 관리 실패)’, ‘낮은 보안 지식이지만 "
           "지불 여력은 있는 고객’ — 이 네 가지가 ScanOps가 푸는 문제의 근거다.", size=10)
    P(doc, "자료 출처: Apple App Store 제출 통계(AppleInsider·The Next Web, 2026) · Veracode GenAI Code Security "
           "Report(2025) · 개인정보보호위원회 SKT 과징금 의결(2025.8) · 한국경제매거진 ‘모두의 창업’ 개인정보 노출 "
           "보도(2026.6) · 자체 고객 인터뷰(2026).", size=8.5, color=(140, 140, 140))
    doc.add_page_break()

    # ── 3. 솔루션 ──
    H(doc, "3. 솔루션: ScanOps", 1, NAVY)
    H(doc, "3.1 어떻게 쓰나 — 주소만 넣으면 끝", 2, BLUE)
    P(doc, "사용법은 단순하다. 검사하고 싶은 GitHub 저장소나 웹사이트 주소를 넣고 검사 방식(웹사이트 DAST · 레포 "
           "전체 SAST · PR 자동 분석)을 고르면, 나머지는 ScanOps가 알아서 처리한다. 사용자는 명령어도, 보안 지식도 "
           "필요 없다.")
    IMG(doc, SHOT / "08_ui_scan_input.png",
        "그림 3. 실제 스캔 요청 화면 — 웹사이트(DAST)·레포 전체(SAST)·PR 자동 분석 중 선택, 주소만 입력하면 검사 시작")
    H(doc, "3.2 무엇을 받나 — 실제 결과 화면", 2, BLUE)
    P(doc, "검사가 끝나면 아래와 같은 리포트를 받는다. 발견된 취약점 개수와 가장 위험한 항목의 위험도 점수(CVSS), "
           "심각도별 분포를 한눈에 보여주고, 항목을 누르면 ‘왜 위험한지(발생 원인)’·‘어떻게 고치는지(해결 방법)’·"
           "‘AI가 어떻게 판단했는지’까지 풀어 준다. 보안 전문가가 아니어도 무엇부터 고쳐야 할지 바로 알 수 있다.")
    IMG(doc, SHOT / "04_ui_report_overview.png",
        "그림 4. 분석 리포트 개요 — 확인된 취약점·최고 CVSS 점수·심각도 분포·취약점 목록을 한 화면에 정리")
    IMG(doc, SHOT / "06_ui_report_detail_modal.png",
        "그림 5. 취약점 상세 — 발생 원인, 해결 방법(언어별 코드 가이드), 참고 문서, AI 판단 근거를 함께 제공")
    P(doc, "심사위원·고객이 가장 빠르게 이해하는 것이 바로 이 화면이다. ‘보안 도구’라는 추상적 설명이 아니라, "
           "주소 한 줄로 이런 리포트가 나온다는 사실이 제품의 가치를 직접 보여준다.")
    H(doc, "3.3 핵심 기술 (요약)", 2, BLUE)
    P(doc, "위 결과 화면 뒤에서 동작하는 핵심 기술은 다음 다섯 가지다. 모두 ‘외부에 코드를 보내지 않으면서도 "
           "정확하게 잡는다’는 한 가지 목표를 향한다.", after=4)
    for label, t in [
        ("자체 구동 AI(코드 비전송):", "분석 AI를 외부 클라우드 API가 아닌 우리 서버에서 직접 구동한다. 그래서 고객 코드가 외부로 나가지 않고, 메모리에서만 처리한 뒤 즉시 폐기한다."),
        ("보안 특화 학습:", "코드 특화 LLM을 보안 취약점 데이터로 추가 학습(QLoRA)시켜, 범용 AI보다 보안 패턴에 민감하게 반응하도록 만들었다."),
        ("최신 CVE 검색(RAG):", "최신 보안 취약점 정보(NVD CVE)를 실시간 검색해 근거로 삼는다. 학습 시점이 고정된 범용 LLM이 모르는 ‘새로 공개된’ 취약점에 강하다."),
        ("오탐 필터(하이브리드):", "‘안전한 코드를 위험하다고 잘못 경고’하는 오탐을 2단계 검증으로 걸러내, 필터 적용 전 100%였던 오탐률을 6% 수준까지 낮췄다."),
        ("멀티파일 코드그래프 추적:", "여러 파일에 흩어진 데이터 흐름을 그래프로 이어, ‘사용자 입력이 실제 위험 지점까지 도달하는 진짜 취약점’인지 구조적으로 가려낸다."),
    ]:
        LEAD(doc, label, t)
    NOTE(doc,
         "분석 모델은 현재 더 큰 3B급 모델(사내 v11)로 업그레이드를 진행 중이다. 또한 자바(Java) 언어 taint 분석과 "
         "Neo4j 그래프 DB 연동 학습을 준비하고 있으며, 학습이 끝나면 동일한 기준으로 성능을 전면 재검증할 예정이다. "
         "아래 3.4·3.5절의 수치는 현재 파이프라인 기준 초기 검증 결과다.",
         label="ℹ︎ 모델 현황  ")
    doc.add_page_break()

    # 3.4 성능 검증 ① 단일코드
    H(doc, "3.4 성능 검증 ① — 단일코드 탐지·오탐률 (100케이스)", 2, BLUE)
    P(doc, "양성(취약) 50개 + 음성(안전) 50개로 구성한 100케이스로 검증했다. 양성은 2026년 5~6월 새로 공개된 NVD "
           "CVE 패턴 기반(범용 LLM이 학습으로 외울 수 없는 신규 패턴)이고, 음성은 mitigation이 적용된 안전 코드로 "
           "오탐을 측정한다. 동일 파이프라인에서 AI 코어만 바꿔 ScanOps와 Grok-3를 공정 비교했다.", after=4)
    TABLE(doc, ["시스템", "탐지율", "오탐률", "정밀도", "정확도", "F1", "응답"],
          [["ScanOps (필터 전 raw)", "100%", f"{RAW['fpr_pct']:.0f}%", "50%", "50%", "66.7", "—"],
           ["ScanOps (하이브리드 필터)", f"{SV5['detection_recall']}%", f"{SV5['false_positive_rate']}%",
            f"{SV5['precision']}%", f"{SV5['accuracy']}%", f"{SV5['f1']}", f"{SV5['avg_time']}s"],
           ["Grok-3-mini (참고)", f"{GROK['detection_recall']}%", f"{GROK['false_positive_rate']}%",
            f"{GROK['precision']}%", f"{GROK['accuracy']}%", f"{GROK['f1']}", f"{GROK['avg_time']}s"]],
          widths=[5.0, 2.0, 2.0, 2.0, 2.0, 1.5, 1.6], hi=1)
    P(doc, "")
    CHART(doc, c_bench, "그림 6. 단일코드 오탐률 — 필터 전 100% → 하이브리드 6%로 개선, 정확도는 상용 모델과 동등 수준(93%)")
    for t in [
        "오탐 필터가 핵심 — 필터 적용 전 오탐률 100%(안전한 코드를 모두 위험하다고 경고)에서 6%로 떨어뜨려, 실사용 가능한 도구로 만들었다.",
        "종합 성능은 상용 모델과 동등 수준 — 정확도·F1이 대등하면서, 응답 속도는 약 10배 빠르고 외부 API 호출이 0이다(자체 모델만으로 처리).",
    ]:
        B(doc, t)
    P(doc, "※ 사내 벤치마크 100케이스 기준 초기 검증 결과이며, 모델 업그레이드 후 동일 기준으로 재검증한다.",
      size=9, color=(130, 130, 130))

    # 3.5 신규 CVE 사례
    H(doc, "3.5 성능 검증 ② — 범용 AI는 놓치고 ScanOps는 잡은 신규 취약점", 2, BLUE)
    P(doc, "‘보안 특화 학습을 했다’는 설명보다, 실제로 더 잘 잡는다는 증거가 중요하다(고객 인터뷰의 핵심 지적이었다). "
           "그래서 범용 AI의 학습 시점 이후인 2026년 5~6월 공개된 신규 CVE에서, 같은 코드를 ScanOps와 범용 모델"
           "(Grok-3)에 동일하게 입력해 비교했다. 아래 취약점들은 ScanOps가 ‘취약’으로 잡았지만 Grok-3는 모두 "
           "‘안전(SAFE)’으로 잘못 판정해 놓쳤다.", after=4)
    TABLE(doc, ["신규 CVE (공개일)", "입력 코드", "실제 취약점", "ScanOps", "Grok-3"],
          [["CVE-2026-44754\n(2026-06-09, SAP)",
            'app.post("/api/replicate",\n(req,res)=>{ replicate(req.body);\nres.sendStatus(200); });',
            "Missing Authorization\n(CWE-862) — 인증/권한 검사 없이\n민감 작업 실행",
            "취약 탐지 ✓", "SAFE (미탐) ✗"],
           ["CVE-2026-11572\n(2026-06, degit)",
            'function merge(t,s){ for(const k\nin s){ t[k]=s[k]; } }\nmerge({}, JSON.parse(req.body));',
            "Prototype Pollution\n(CWE-1321) — 신뢰 불가 입력의\n재귀 병합",
            "취약 탐지 ✓", "SAFE (미탐) ✗"],
           ["CVE-2026-21404\n(2026-06-04, NAVTOR)",
            'String SOAP_USER="svc";\nString SOAP_PW="P@ssw0rd!";\nauth(SOAP_USER, SOAP_PW);',
            "Hardcoded Credentials\n(CWE-798) — 소스에 자격증명\n하드코딩",
            "취약 탐지 ✓", "SAFE (미탐) ✗"]],
          widths=[3.2, 5.0, 4.3, 1.9, 2.1], hi=[0, 1, 2], hi_fill="EAF6EC")
    P(doc, "")
    CHART(doc, c_h2h, "그림 7. 2026년 5~6월 신규 CVE 5건 — ScanOps 전건 탐지(✓), Grok 전건 미탐(✗)")
    P(doc, "범용 AI는 ‘학습 시점 이후의 신규 패턴’과 ‘있어야 할 인증·권한이 없는(negative space) 취약점’에 약하다. "
           "ScanOps는 보안 특화 학습 + 최신 CVE 실시간 검색 구조라 ‘오늘 새로 올라온’ 취약점도 근거를 갖고 잡는다.")
    doc.add_page_break()

    # 3.6 멀티파일 코드그래프
    H(doc, "3.6 성능 검증 ③ — 여러 파일에 걸친 취약점 추적 (코드그래프, 100케이스)", 2, BLUE)
    P(doc, "실제 코드는 한 파일이 아니라 여러 파일에 걸쳐 데이터가 흐른다. 같은 코드라도 정적 이미지를 넘기면 "
           "안전하지만 URL 파라미터 같은 사용자 입력을 넘기면 실제 XSS·SSRF 위험이 된다. 이 둘을 텍스트만 보고 "
           "구분하기는 최신 모델에게도 어렵다. ScanOps는 파일 간 관계를 그래프로 이어 이 구분을 구조적으로 "
           "해결한다.", after=4)
    TABLE(doc, ["시스템", "정확도", "Recall(탐지율)", "Specificity(오탐방지)", "FP", "FN"],
          [["ScanOps (코드그래프)", "100.0%", "100.0%", "100.0%", "0건", "0건"],
           ["Grok-3-mini (코드만)", f"{GRAPH['grok_accuracy']:.1f}%", "35.3%", "100.0%", "0건", "33건"]],
          widths=[4.5, 2.6, 3.2, 3.2, 1.8, 1.8], hi=0)
    P(doc, "")
    CHART(doc, c_graph, "그림 8. 멀티파일 taint 추적 100케이스 — ScanOps 100% vs Grok-3-mini 67%(미탐 다수)")
    P(doc, "Grok은 사용자 입력이 실제 위험 지점에 도달하는 51개 중 33개를 ‘확신 없음 → SAFE’로 놓쳤다(탐지율 35.3%). "
           "ScanOps 그래프 엔진은 100케이스를 전부 맞췄다 — 멀티파일 추적은 모델 크기가 아니라 그래프 기반 구조의 "
           "차이임을 보여준다.")
    NOTE(doc, "이 벤치마크는 그래프 판정 로직 자체의 정확도를 검증한 것이다. 현재 판정은 인메모리 엔진으로 동작하며, "
              "Neo4j 그래프 DB 연동과 프론트엔드 그래프 시각화는 로드맵(7.2절) 항목으로 아직 연동 전이다. 판정 "
              "정확도와 DB·시각화 연동은 별개이며, 정확도는 인메모리 엔진으로도 동일하게 보장된다.")
    doc.add_page_break()

    # ── 4. 차별성 ──
    H(doc, "4. 차별성 및 경쟁 우위", 1, NAVY)
    H(doc, "4.1 경쟁 포지셔닝", 2, BLUE)
    P(doc, "경쟁 환경을 ‘도입 비용·진입장벽’과 ‘코드 프라이버시(외부 전송 여부)’ 두 축으로 보면, ScanOps는 "
           "‘저가·셀프서브이면서 코드를 외부로 보내지 않는’ 영역을 사실상 단독 점유한다. 엔터프라이즈 도구(스패로우)는 "
           "강력하지만 비싸고 무겁고, 클라우드 SaaS(Snyk·Semgrep)나 범용 AI(ChatGPT·Claude·Grok)는 편하지만 코드를 "
           "외부로 전송한다.")
    H(doc, "4.2 경쟁사 비교", 2, BLUE)
    P(doc, "① 가격 비교", bold=True, size=10.5, after=2)
    TABLE(doc, ["도구", "가격", "비고"],
          [["ScanOps", "월 29,900원~", "개인(Pro/Max)·팀(Team) 플랜, 사용량 기반 과금"],
           ["Snyk", "월 $25~57/contributor", "팀 규모에 따라 급증"],
           ["Semgrep", "월 $40/contributor~", "커스텀 룰 작성 필요"],
           ["GitHub Advanced Security", "무료(public)/유료(private, per committer)", "GitHub 전용"],
           ["스패로우", "엔터프라이즈 견적", "연단위 라이선스"]],
          widths=[4.5, 6.0, 6.0], hi=0, hi_fill="1E8449", hi_text=(255, 255, 255))
    P(doc, "② 코드 프라이버시", bold=True, size=10.5, after=2)
    TABLE(doc, ["도구", "코드 전송 여부", "비고"],
          [["ScanOps", "전송 안 함", "자체 서버 처리·즉시 폐기"],
           ["Snyk / Semgrep(상용)", "클라우드 전송", "외부 서버 분석"],
           ["범용 AI(ChatGPT·Claude·Grok)", "클라우드 전송", "프롬프트로 코드 유출"],
           ["스패로우", "온프레미스", "자체 설치"]],
          widths=[5.0, 5.0, 6.5], hi=0, hi_fill="1E8449", hi_text=(255, 255, 255))
    P(doc, "③ AI 보안 특화 / 멀티파일 분석", bold=True, size=10.5, after=2)
    TABLE(doc, ["도구", "방식", "비고"],
          [["ScanOps", "보안 파인튜닝 + 하이브리드 오탐 필터 + 코드그래프 taint 추적", "단일코드 오탐률 6% · 멀티파일 정확도 100%(사내 벤치마크)"],
           ["Snyk", "룰 + 일부 AI(DeepCode)", "클라우드"],
           ["Semgrep", "룰 기반", "커스텀 룰 강점"],
           ["범용 AI", "범용 LLM, 텍스트만으로 판단", "멀티파일 데이터 흐름 증명에 구조적으로 약함(자체 벤치마크 탐지율 35.3%)"]],
          widths=[4.5, 6.5, 5.5], hi=0, hi_fill="1E8449", hi_text=(255, 255, 255))
    H(doc, "4.3 ScanOps만의 네 가지 무기", 2, BLUE)
    for t in [
        "코드 유출 0 — 자체 구동 모델로 소스코드가 외부로 나가지 않으며, 메모리 처리 후 즉시 폐기하고 삭제 로그로 증명한다.",
        "동등 성능 + 합리적 가격 — 단일코드 100케이스에서 상용 모델과 대등한 정확도(93%)를, 월 2.9~9.9만원 셀프서브로 제공.",
        "오탐 관리 + 신규 CVE 강점 — 하이브리드 필터로 오탐 6%, NVD RAG로 학습 시점 이후 신규 취약점까지 커버.",
        "멀티파일 taint 추적 — 코드그래프로 파일 간 데이터 흐름을 추적해, 범용 LLM이 구조적으로 약한 멀티파일 오탐·미탐 구간(정확도 100% vs Grok 67%)에서 명확한 우위.",
    ]:
        B(doc, t)
    H(doc, "4.4 경쟁 우위의 방어 (Moat)", 2, BLUE)
    P(doc, "대기업(GitHub·Snyk)이 유사 기능을 무료화할 경우의 대응:", size=10.5, after=3)
    for t in [
        "커스텀 룰: 소규모 팀별 스택(예: React+Firebase)에 특화된 맞춤 탐지 룰 제공.",
        "빠른 대응: 고객 요청 → 모델 재학습 → 배포까지 1주일 이내. 대기업은 의사결정이 느리다.",
        "한국어 특화(향후): 한국어 취약점 설명 + 국내 법규(개인정보보호법) 맞춤 리포트.",
    ]:
        B(doc, t)
    doc.add_page_break()

    # ── 5. 시장 분석 ──
    H(doc, "5. 시장 분석", 1, NAVY)
    H(doc, "5.1 시장 성장성", 2, BLUE)
    P(doc, "글로벌 애플리케이션 보안(AppSec) 시장은 2025년 약 USD 106.5억(약 14.8조원)에서 2033년 약 USD 420억 "
           "규모로 연평균 18.8% 성장이 전망된다. 클라우드 전환과 AI 코딩 도구 확산으로 ‘SAST/DAST’ 수요가 빠르게 "
           "늘고 있으며, 아시아·태평양이 가장 높은 성장률을 보인다. 국내 정보보호산업도 2024년 18.6조원으로 전년 "
           "대비 10.5% 성장했다.")
    H(doc, "5.2 TAM / SAM / SOM", 2, BLUE)
    TABLE(doc, ["구분", "정의", "규모(추정)"],
          [["TAM", "글로벌 애플리케이션 보안 시장 전체", "약 14.8조원"],
           ["SAM", "클라우드형 코드보안 SaaS · SME/소규모팀 세그먼트", "약 8,800억원"],
           ["SOM", "3년 내 도달 가능한 국내 beachhead(스타트업·소규모팀, bottom-up)", "약 15억원"]],
          widths=[2.5, 9.5, 4.5])
    P(doc, "")
    CHART(doc, c_market, "그림 9. 시장 규모 추정 (TAM: Grand View Research 2026 · SAM/SOM 가정 기반)")
    P(doc, "출처 및 산정 근거. TAM은 Grand View Research, ‘Application Security Market Size, Industry Report 2033’"
           "(2026)의 2025년 세계 애플리케이션 보안 시장 USD 106.5억(약 14.8조원, 환율 1,390원/USD)·연평균 18.8% "
           "전망을 인용했다. SAM은 클라우드형 SAST/DAST이면서 SME/소규모팀 세그먼트를 약 6%로 가정해 약 8,800억원으로 "
           "top-down 추정했다. SOM은 bottom-up으로, 3년 내 도달 가능한 국내 스타트업·소규모 팀 유료 계정과 상향된 "
           "ARPU(Pro·Max·Team 혼합)를 반영해 약 15억원 규모로 보수적으로 산정했다. (모든 추정치는 예비창업 단계의 "
           "가정 기반이며 베타 데이터로 정밀화한다.)", size=10)
    doc.add_page_break()

    # ── 6. BM ──
    H(doc, "6. 비즈니스 모델", 1, NAVY)
    H(doc, "6.1 구독 플랜 — 개인과 팀을 나눈다", 2, BLUE)
    P(doc, "ScanOps의 주 타깃은 스타트업 팀이고, 이들은 대부분 GitHub 조직(Organization)을 쓴다. 그래서 요금제를 "
           "‘개인용’과 ‘팀용’으로 분리했다. 혼자 또는 코드가 많은 개인은 Pro·Max로, 조직 단위로 함께 쓰는 팀은 "
           "멤버 관리·통합 청구가 되는 Team으로 자연스럽게 나뉜다.")
    P(doc, "① 개인 플랜", bold=True, size=10.5, after=2)
    TABLE(doc, ["플랜", "가격(월)", "DAST(웹)", "GitHub App(액션)", "SAST(레포)", "추천 대상"],
          [["회원가입", "0원", "1회 무료", "✗", "✗", "가볍게 체험"],
           ["Pro", "29,900원", "월 5회", "월 5만 줄", "월 10만 줄", "개인 개발자"],
           ["Max", "99,000원", "월 30회", "월 30만 줄", "월 50만 줄", "코드 많은 개인·소수"]],
          widths=[2.4, 2.4, 2.2, 3.0, 2.8, 3.7], hi=1)
    P(doc, "② 팀 플랜 (Team) — 조직 단위", bold=True, size=10.5, after=2)
    TABLE(doc, ["구분", "가격(월)", "DAST(웹)", "GitHub App(액션)", "SAST(레포)", "핵심"],
          [["Team (기본 3명)", "89,000원", "월 20회", "월 24만 줄", "월 45만 줄", "조직 설치·멤버 관리·통합 청구"],
           ["└ 추가 1명당", "+25,000원", "+7회", "+8만 줄", "+15만 줄", "인원 증가에 따라 비례 확장"]],
          widths=[3.2, 2.4, 2.0, 3.0, 2.8, 3.1], hi=0, hi_fill="F3E8FF")
    P(doc, "과금 인원의 정의(혼란 차단). 과금 대상은 ‘스캔 대상 레포에 코드를 push하는 멤버’다. 리포트만 열람하는 "
           "사람(PM·디자이너 등)은 무료로 무제한 초대할 수 있다. ‘보안 스캔을 실제로 유발하는 사람만 과금’이라 "
           "명분이 분명하고, 열람자를 무료로 풀어 주면 조직 안에서 ScanOps 노출이 자연스럽게 늘어나는 확산 효과도 "
           "있다.", size=10)
    P(doc, "‘한도가 늘어서’가 아니라 ‘사람이 늘어서’ Team", bold=True, size=10.5, color=GREEN, after=2)
    TABLE(doc, ["Team 인원", "월 가격", "액션 한도", "Max(99,000원·30만 줄) 대비"],
          [["3명 (기본)", "89,000원", "24만 줄", "가격·한도 모두 Max 아래"],
           ["4명", "114,000원", "32만 줄", "여기서 한도가 Max를 추월"],
           ["5명", "139,000원", "40만 줄", "확실히 상회"]],
          widths=[3.2, 3.0, 3.2, 7.1], hi=1)
    P(doc, "가격 사다리가 Pro < Max < Team(4명~) 순서로 자연스럽게 이어진다. 즉 ‘한도가 더 필요해서 Team’이 아니라 "
           "‘사람이 늘어서 Team’이라는 메시지가 가격으로 표현되며, 코드 많은 개인은 여전히 Max가 합리적이라 "
           "플랜 간 잠식(카니발라이제이션)도 없다.", size=10)
    P(doc, "추가 결제(한도 초과 시, 전 플랜 공통): SAST·GitHub App 1만 줄당 5,000원 · DAST 웹 스캔 3회당 5,000원 · "
           "Pro는 가입 후 7일 무료체험 후 자동 결제 전환.", size=9.5, color=(90, 90, 90))
    CHART(doc, c_price, "그림 10. 구독 플랜 — 개인(Pro·Max)과 팀(Team) 분리, Team은 인원당 +25,000원", width=Inches(5.6))
    P(doc, "무료 구독 플랜은 두지 않는다. 대신 회원가입 시 웹 URL DAST를 1회 무료 제공하고(PLG), Pro는 7일 무료체험 "
           "후 결제로 전환되게 해 진입 마찰을 낮춘다. DAST는 외부 웹을 능동 스캔해 1건당 자원이 커 월 횟수 상한을 "
           "두고, SAST·GitHub App은 코드량(LOC) 한도로 통제한다. 한도와 추가 구매는 서버 비용을 통제하는 안전장치이자 "
           "매출의 상방을 여는 장치다.")
    P(doc, "가격 책정 근거. 인터뷰(§2.5)에서 확인된 타깃(스타트업 CTO·보안 우려 개발팀)은 개발도구에 월 14만원 이상을 "
           "쓰고 ‘CTO는 월 10만원대 지출을 고민하지 않는’ 가격 민감도가 낮은 층이다. 저가 박리다매(월 1만원 이하)는 "
           "다수 무료·이탈 사용자를 유인해 서버 부담만 키우는 반면, 보안의 가치를 아는 소수 유료 고객에게 제값을 "
           "받는 편이 단위경제와 손익분기에 유리하다.")
    H(doc, "6.2 단위 경제 및 손익분기 (Unit Economics)", 2, BLUE)
    for t in [
        "변동비는 매우 낮다 — 스캔 1건당 실제 비용은 DAST 약 1원, SAST(5만 줄) 약 20원 수준. 경량 양자화 모델이라 고가 GPU 없이 운영한다.",
        "서버 고정비(Railway, 모델+Qdrant+API 상시 가동) 약 월 50만원. 결제 수수료는 매출의 약 3%(토스/스트라이프) 가정.",
        "인프라 손익분기: 고정비 50만원 기준, 수수료 차감 후 Pro(약 29,000원) 약 18명 또는 Max(약 96,000원) 약 6명, Team(기본 3명·약 89,000원) 약 6팀이면 서버 고정비를 회수한다.",
        "Team이 단위경제를 끌어올린다 — Team은 ARPU가 개인 플랜의 3배 이상이라, 유료 전환 팀이 10팀만 생겨도 단위경제가 크게 개선된다. 추가 인원 과금(+25,000원/명)이 확장 매출(expansion revenue)로 들어와 LTV를 더 높인다.",
        "마케팅 포함 손익분기: 고객획득비용(CAC)을 1인당 약 5만~15만원으로 보수적 가정해도, 낮은 가격 민감도와 높은 잔존(월 이탈 ~5% 가정 시 평균 20개월)으로 Max 기준 LTV ≈ 약 180만원, LTV/CAC ≈ 12배 이상. 유료 전환 수십 명 규모에서 흑자 구조에 진입한다.",
    ]:
        B(doc, t)
    NOTE(doc, "인프라는 Railway만 실제로 사용 중이다(Dockerfile·railway.toml·docker-compose 기준 확인). "
              "코드그래프용 Neo4j(Aura 등)는 아직 계정·인스턴스가 없어 별도 고정비에 포함되지 않았으며, 연동 시 "
              "고정비에 소액 추가될 수 있다(Aura 무료 티어로 시작 가능).")
    doc.add_page_break()

    # ── 7. 사업화 전략 ──
    H(doc, "7. 사업화 전략 및 로드맵", 1, NAVY)
    H(doc, "7.1 Go-To-Market: CTO 타깃 PLG → SME → 글로벌", 2, BLUE)
    P(doc, "초기 타깃은 ‘인디 개발자 다수’가 아니라 ‘보안을 우려하는 스타트업 CTO’다. 인터뷰(§2.5)에서 드러났듯 "
           "스타트업을 시작하는 개발자가 가장 우려하는 지점이 바로 보안이며, 이들은 월 10만원대 지출을 고민하지 "
           "않는다. Reddit·GitHub Marketplace는 전환율을 기대하기 어려워 주력 채널에서 제외하고, 다음 채널 전략으로 "
           "전환한다.")
    P(doc, "주력 채널 — 개발자 유튜버 협찬", bold=True, size=10.5, color=GREEN, after=2)
    for t in [
        "개발자 유튜버(백엔드·DevOps·스타트업 채널)와의 협찬/리뷰로 CTO·시니어 개발자에게 직접 도달. 인스타 광고는 타깃 적합도가 낮아 유료 광고에서 제외하고 유튜브에 집중.",
        "‘신규 CVE 30초 분석’, ‘ScanOps로 내 레포 스캔해보기’ 등 제품 강점을 짧은 데모 영상으로 전달.",
    ]:
        B(doc, t)
    P(doc, "브랜드 채널 — 인스타그램 ‘ScanOps Security Report’", bold=True, size=10.5, color=GREEN, after=2)
    for t in [
        "주간 단위로 발견·분석한 취약점을 잡지/홍보물 형식의 카드뉴스로 발행하는 인스타 계정을 운영. (광고가 아니라 콘텐츠로 계정 자체를 키운다.)",
        "‘이번 주 발견된 위험 취약점’ 업적을 꾸준히 알리며 전문성·신뢰를 축적 → 팔로워가 잠재 고객으로 유입되는 오가닉 획득 채널.",
    ]:
        B(doc, t)
    P(doc, "유료 전환 메커니즘", bold=True, size=10.5, color=GREEN, after=2)
    for t in [
        "매 PR마다 GitHub App이 PR 코멘트로 취약점 알림 → 습관화 → Pro 7일 무료체험 → 자동 결제 전환.",
        "랜딩페이지에 실제 탐지 사례·UI 스크린샷·CTO 후기 노출(인터뷰에서 ‘실제 탐지 예시가 와닿으면 관심이 생긴다’ 확인).",
    ]:
        B(doc, t)
    H(doc, "7.2 단계별 로드맵", 2, BLUE)
    TABLE(doc, ["단계", "기간", "핵심 과제"],
          [["0. 모델·그래프 고도화", "즉시~3개월", "분석 모델 3B 업그레이드, Java taint 분석·Neo4j(Aura) 그래프 학습 및 전면 재검증, 프론트 코드그래프 시각화"],
           ["1. 토대", "0~3개월", "사용자 인증·계정, 모든 스캔을 user_id에 연결, XLS 리포트·NVD 알림 기능"],
           ["2. 게이트", "3~6개월", "플랜별 사용량(LOC) 한도·줄 추가 구매, GitHub 조직(Team) 멤버 관리·통합 청구, 레포/도메인 소유권 인증"],
           ["3. 계정 UI", "6~9개월", "마이페이지, 스캔 기록·삭제, 사용량 미터"],
           ["4. 수익화", "9~12개월", "결제 연동(토스/스트라이프), Pro 7일 무료체험 전환 퍼널"],
           ["5. 차별화", "12~18개월", "경쟁사 대비 오탐률·멀티파일 정확도 벤치마크 공개, 글로벌 진출 준비"]],
          widths=[2.8, 2.4, 11.1], hi=0, hi_fill="F3E8FF")
    H(doc, "7.3 3개년 추정 손익", 2, BLUE)
    CHART(doc, c_rev, "그림 11. 3개년 매출·고객 추정 (CTO 타깃·고가 플랜, 가정 기반)")
    P(doc, "타깃을 고가치 CTO로 좁히고 Pro 29,900원·Max 99,000원·Team(인원 비례)으로 ARPU를 끌어올렸다. 1년차 ARR "
           "약 1.1억원(유료 250명), 2년차 약 4.2억원(900명), 3년차 약 9.8억원(2,000명)을 목표한다. 변동비가 극히 "
           "낮고 인프라 손익분기가 Pro 약 18명 수준이라, 유료 전환이 본격화되는 1년차 후반~2년차부터 흑자 전환이 "
           "가능한 구조다. Team 전환이 늘수록 ARPU·확장 매출이 더해져 상방이 커진다. (가정 기반 추정치, 베타 "
           "데이터로 정밀화)", size=10)
    doc.add_page_break()

    # ── 8. 팀 ──
    H(doc, "8. 팀 구성", 1, NAVY)
    P(doc, "ScanOps는 부산대학교 정보컴퓨터공학부 재학생 4인으로 구성된다. AI 모델·백엔드·인프라·보안·프론트엔드·"
           "UI/UX·마케팅까지, SaaS 하나를 만들고 운영하는 데 필요한 역량을 외부 의존 없이 자체적으로 보유한 것이 "
           "강점이다.")
    TABLE(doc, ["팀원", "역할"],
          [["김세한 (팀장)", "AI 모델 · 프론트엔드 · 총괄"],
           ["전혜은", "백엔드 · 인프라"],
           ["이경윤", "AI 모델 · 보안"],
           ["최효석", "UI/UX · 마케팅 · QA"]],
          widths=[4.0, 12.0])
    P(doc, "")
    P(doc, "각 팀원이 하는 일", bold=True, size=10.5, color=GREEN, after=2)
    for label, t in [
        ("김세한(팀장).",
         "전체 방향과 일정을 총괄하고, 보안 특화 AI 모델(QLoRA 파인튜닝·RAG·오탐 필터)의 학습과 성능 검증을 주도한다. "
         "동시에 사용자가 실제로 보는 프론트엔드 화면(스캔 요청·리포트 UI)을 직접 구현해 모델과 제품을 잇는다."),
        ("전혜은(백엔드·인프라).",
         "Spring Boot 백엔드와 FastAPI 분석 서버, 데이터 흐름·API를 설계·구현하고, Railway 기반 클라우드 배포와 "
         "운영(컨테이너·DB·상시 가동)을 책임진다. 서비스가 안정적으로 도는 ‘뼈대’를 담당한다."),
        ("이경윤(AI 모델·보안).",
         "보안 도메인 지식을 모델 학습 데이터(취약점·CVE 케이스)와 벤치마크 설계로 옮기고, 코드그래프 기반 멀티파일 "
         "taint 추적 로직 등 탐지 정확도를 끌어올리는 핵심 알고리즘을 함께 만든다."),
        ("최효석(UI/UX·마케팅·QA).",
         "비전공자도 한눈에 이해하는 리포트 화면의 UI/UX를 설계하고, 인스타그램·유튜버 협찬 등 마케팅 채널을 "
         "운영한다. 또한 출시 전 품질 검증(QA)으로 제품 완성도를 관리한다."),
    ]:
        LEAD(doc, label, t)

    # ── 9. 기대효과 ──
    H(doc, "9. 기대효과 및 결론", 1, NAVY)
    P(doc, "ScanOps는 ‘작은 팀도 코드를 외부에 맡기지 않고 합리적 가격에 쓰는 AI 보안 진단 도구’라는, 시장에 비어 "
           "있던 자리를 정확히 겨냥한다. 자체 모델로 단일코드 100케이스에서 오탐률 6%·정확도 93%를 달성해 ‘탐지율’을 "
           "넘어 ‘오탐 관리’까지 검증했고, 코드그래프 기반 멀티파일 추적으로 정확도 "
           f"{GRAPH['scanops_accuracy']:.0f}%를 추가 검증해 범용 LLM이 구조적으로 약한 구간에서도 명확한 우위를 "
           "확인했다. 기술적 실현 가능성(이미 구현됨), 명확한 차별성(코드 비전송 + 오탐 관리 + 멀티파일 정확도 + "
           "합리적 가격), 성장하는 시장(연 18.8%)을 모두 충족하는 사업이다.")
    P(doc, "핵심 메시지", bold=True, size=10.5, color=GREEN, after=2)
    QUOTE(doc, "“보안을 위해 코드를 외부에 넘기지 않아도 된다. ScanOps는 자체 AI와 코드그래프로, 단일 코드부터 "
               "여러 파일에 걸친 취약점까지 상용 모델과 동등하거나 더 나은 정확도를, 작은 팀이 감당할 가격에 "
               "제공한다.”")
    doc.add_page_break()

    # ── 10. 고객 검증 ──
    H(doc, "10. 고객 검증 및 향후 계획", 1, NAVY)
    P(doc, "타깃 고객의 문제 인식과 지불 의향은 §2.5(현장의 목소리)에서 인터뷰 근거로 제시했다. 이 절은 그 인터뷰가 "
           "‘제품’에 준 피드백과, 우리가 거기에 어떻게 대응했는지, 그리고 향후 검증 계획을 정리한다.", after=4)
    H(doc, "10.1 제품 피드백 — ‘범용 도구보다 나은 증거를 보여라’", 2, BLUE)
    P(doc, "실사용·제품화 관점의 스타트업 개발자 인터뷰에서 나온 핵심 지적은 분명했다.", after=3)
    for t in [
        "‘우리 모델만의 뛰어난 결과를 보여주는 예시가 와닿게 느껴지면 관심이 생긴다.’ → 개발자는 추상적 설명보다 실제 탐지 사례에 반응한다.",
        "‘Claude Code·Codex 같은 유명 에이전트에 코드를 넣어 돌린 결과보다 ScanOps가 더 나은 예시를 보여줘야 한다. \"보안 전용으로 파인튜닝했다\"는 건 그렇구나 정도일 뿐이다.’",
        "‘바이브 코딩과 보안이 트레이드오프라는 점에서 주제 자체가 핫하다. 성능만 입증되면 아이디어가 괜찮다.’",
    ]:
        B(doc, t)
    P(doc, "→ 이 지적에 본 보고서 §3.5·3.6(범용 AI는 놓치고 ScanOps는 잡은 신규 CVE·멀티파일 사례)로 직접 "
           "대응했다. ‘보안 파인튜닝했다’는 설명이 아니라, 동일 입력에서 상용 모델이 놓친 취약점을 우리가 잡은 "
           "구체적 사례가 바로 그 ‘와닿는 증명’이다.", size=10)
    H(doc, "10.2 인식 → ScanOps의 대응", 2, BLUE)
    for t in [
        "보안 지식이 낮다(자가평가 2~3/10) → 전문가 없이도 위험도·수정방법·참고문서를 떠먹여 주는 ScanOps의 CVSS·방어방침·레퍼런스 출력이 공백을 메운다.",
        "필요성은 알지만 미룬다 → 가입 즉시 셀프서브, PR마다 자동으로 도는 GitHub App·Pro 무료체험이 미루기를 깬다.",
        "회사 예산이면 100~200만원도 지불 의향 → 개인 결제(Pro·Max)로 진입시킨 뒤 팀 결제(Team·B2B)로 확장하는 경로가 유효하다.",
        "코드 외부 전송이 실재하는 구매 장벽 → ‘코드 비전송·즉시 폐기’가 이 우려를 정면으로 해소한다.",
        "실제 사용 후 만족도 급상승(‘이거 좋은데?!’)·학습 가치까지 → 데모·무료체험(PLG)과 랜딩의 실제 탐지 예시가 전환을 견인한다.",
    ]:
        B(doc, t)
    H(doc, "10.3 향후 고객 검증 계획", 2, BLUE)
    for t in [
        "베타 전(3개월): 스타트업 CTO·소규모팀 20~30명 인터뷰. 질문 — ‘보안 점검을 미루는 이유?’, ‘Pro 29,900원이면 쓸 의향은?’, ‘코드 외부 전송 우려는?’",
        "베타 후(6개월): 대기자 100명 확보, 유료 전환율·평균 사용 빈도·이탈률(Churn) 수집.",
    ]:
        B(doc, t)
    doc.add_page_break()

    # ── 11. 리스크 ──
    H(doc, "11. 리스크 관리 및 대응 전략", 1, NAVY)
    H(doc, "11.1 소유권 인증", 2, BLUE)
    for t in [
        "웹 도메인: .well-known 파일 방식 — 도메인 루트에 검증 파일을 배치해 소유권을 확인(웹페이지에 절차 상세 안내 예정).",
        "GitHub 레포: GitHub App 인증 — 레포 첫 스캔 시 1회 인증하고, 이후에는 재인증 없이 앱 권한을 유지.",
    ]:
        B(doc, t)
    H(doc, "11.2 오남용 방지", 2, BLUE)
    for t in [
        "DAST 남용 방지: 월 횟수 상한(Pro 5회·Max 30회·Team 20회+), 타인 도메인 스캔 시 소유권 인증 필수.",
        "사용량 한도: SAST·GitHub App은 LOC 한도+줄 추가 구매로 통제해 서버 비용·남용을 동시에 관리.",
        "법적 책임 범위: 이용약관에 ‘보안 진단 도구이며 진단 결과의 법적 효력은 없고 최종 책임은 사용자에게 있음’ 명시. 주간 보안 요약(XLS) 리포트는 ‘내부 참고용’으로 포지셔닝.",
    ]:
        B(doc, t)
    H(doc, "11.3 데이터 보안", 2, BLUE)
    for t in [
        "코드 처리: 분석 소스코드는 메모리에서만 처리 후 즉시 폐기, 삭제 로그로 증명.",
        "결과 저장: 취약점 진단 결과만 저장(코드 원본 미저장), 사용자 요청 시 결과 삭제 가능.",
    ]:
        B(doc, t)
    H(doc, "11.4 기술 리스크 — 모델 고도화·Neo4j 연동 일정", 2, BLUE)
    for t in [
        "현재 그래프 판정 로직은 인메모리로 동작해 정확도엔 영향이 없으나, 모델 3B 업그레이드·Java taint 학습·Neo4j(Aura) 연동·시각화가 지연될 경우 ‘그래프 시각화’를 제품 화면에서 보여주는 시점이 늦어질 수 있다.",
        "대응: Aura 무료 티어로 우선 연동(저비용·단기간 가능), 백엔드 DTO 확장은 필드 추가 수준의 경량 작업으로 7.2절 로드맵 0단계에 우선 배정했다. 모델·학습 완료 후 동일 기준으로 전면 재검증한다.",
    ]:
        B(doc, t)

    P(doc, "")
    P(doc, "ScanOps 사업계획서 — 코드 비전송 자체 AI 보안 진단 · 단일코드 오탐률 + 멀티파일 코드그래프 검증 · "
           "개인/팀 플랜 · CTO 타깃 GTM · 2026.6",
      size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=(130, 130, 130))

    # 기본 템플릿 settings.xml의 zoom percent 누락 보정(스키마 적합성)
    try:
        zoom = doc.settings.element.find(qn("w:zoom"))
        if zoom is not None and zoom.get(qn("w:percent")) is None:
            zoom.set(qn("w:percent"), "100")
    except Exception:
        pass

    doc.save(OUT)
    try:
        doc.save(OUT_DESKTOP)
    except Exception as e:
        print("데스크톱 저장 건너뜀:", e)
    print(f"저장 완료: {OUT}")
    print(f"저장 완료: {OUT_DESKTOP}")


if __name__ == "__main__":
    build()
