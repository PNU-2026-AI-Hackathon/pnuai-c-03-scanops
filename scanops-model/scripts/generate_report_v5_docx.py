"""
ScanOps v5 보고서 DOCX 생성 — 오탐률 중심 (v4 스타일 계승 + 신규 차트)
실행: python scripts/generate_report_v5_docx.py
"""
import io
import json
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

plt.rcParams['font.family'] = ['AppleGothic', 'NanumGothic', 'Malgun Gothic', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

BASE = Path(__file__).resolve().parents[1]
OUT = BASE / "reports" / "ScanOps_졸업과제보고서_v5.docx"
RES = json.loads((BASE / "reports" / "results_v5_false_positive_benchmark.json").read_text())

SV5 = RES["systems"][0]["metrics"]
GROK = RES["systems"][1]["metrics"]
RAW = RES["v4_raw_safe_sample"]


# ── 스타일 헬퍼 (v4 계승) ──────────────────────────────────────────────────
def set_font(run, name="맑은 고딕", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), name)
    rPr.insert(0, rFonts)


def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "맑은 고딕"
        run.font.size = Pt([0, 18, 15, 13][level])
        if color:
            run.font.color.rgb = RGBColor(*color)
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        rPr.insert(0, rFonts)
    return p


def add_para(doc, text="", bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT,
             color=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        set_font(p.add_run(text), size=size, bold=bold, color=color)
    return p


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "Consolas")
    rPr.insert(0, rFonts)
    return p


def add_table(doc, headers, rows, col_widths=None, highlight_row=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(h), size=10, bold=True, color=(255, 255, 255))
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "2C3E50")
        tcPr.append(shd)
    for ri, row_data in enumerate(rows):
        row = table.add_row()
        fill = None
        if highlight_row is not None and ri == highlight_row:
            fill = "EAF6EC"
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(p.add_run(str(val)), size=10, bold=(fill is not None))
            if fill:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
                tcPr.append(shd)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def fig_to_stream(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)
    buf.seek(0)
    return buf


def add_chart(doc, buf, width=Inches(5.9), caption=""):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(buf, width=width)
    if caption:
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption); r.font.size = Pt(9); r.font.italic = True
        r.font.color.rgb = RGBColor(100, 100, 100)


# ── 차트 ────────────────────────────────────────────────────────────────
def chart_fpr():
    fig, ax = plt.subplots(figsize=(9, 3.4), facecolor='#FAFAFA')
    ax.set_facecolor('#F7F9FC')
    names = ['ScanOps v4-raw\n(오탐 필터 없음)', 'ScanOps v5\n(하이브리드)', 'Grok-3-mini']
    vals = [RAW["fpr_pct"], SV5["false_positive_rate"], GROK["false_positive_rate"]]
    colors = ['#E74C3C', '#27AE60', '#3498DB']
    bars = ax.barh(names, vals, color=colors, height=0.55, zorder=3)
    for b, v in zip(bars, vals):
        ax.text(v + 1.5, b.get_y() + b.get_height()/2, f"{v}%", va='center', fontsize=11, fontweight='bold')
    ax.set_xlim(0, 110)
    ax.set_xlabel("오탐률 (False Positive Rate, %) — 낮을수록 좋음", fontsize=10)
    ax.set_title("오탐률 비교: v4-raw → v5 하이브리드 오탐 필터로 94%p 개선", fontsize=12, fontweight='bold', pad=10)
    ax.grid(True, axis='x', linestyle='--', alpha=0.4)
    ax.invert_yaxis()
    fig.tight_layout()
    return fig_to_stream(fig)


def chart_metrics():
    fig, ax = plt.subplots(figsize=(9, 4.6), facecolor='#FAFAFA')
    ax.set_facecolor('#F7F9FC')
    labels = ['탐지율\n(Recall)', '오탐률\n(FPR)', '정밀도\n(Precision)', '정확도\n(Accuracy)', 'F1']
    sv5 = [SV5["detection_recall"], SV5["false_positive_rate"], SV5["precision"], SV5["accuracy"], SV5["f1"]]
    grok = [GROK["detection_recall"], GROK["false_positive_rate"], GROK["precision"], GROK["accuracy"], GROK["f1"]]
    x = np.arange(len(labels)); w = 0.36
    b1 = ax.bar(x - w/2, sv5, w, label='ScanOps v5', color='#27AE60', zorder=3)
    b2 = ax.bar(x + w/2, grok, w, label='Grok-3-mini', color='#3498DB', zorder=3)
    for bars in (b1, b2):
        for b in bars:
            ax.text(b.get_x() + b.get_width()/2, b.get_height() + 1.5, f"{b.get_height():.0f}",
                    ha='center', fontsize=9)
    ax.set_ylim(0, 112)
    ax.set_xticks(x); ax.set_xticklabels(labels, fontsize=10)
    ax.set_ylabel("점수 (%)", fontsize=10)
    ax.set_title("ScanOps v5 vs Grok-3 — 동일 파이프라인, LLM만 교체", fontsize=12, fontweight='bold', pad=10)
    ax.legend(fontsize=10, loc='lower center', ncol=2)
    ax.grid(True, axis='y', linestyle='--', alpha=0.4)
    fig.tight_layout()
    return fig_to_stream(fig)


def chart_pricing():
    fig, ax = plt.subplots(figsize=(8.5, 3.6), facecolor='#FAFAFA')
    ax.set_facecolor('#F7F9FC')
    plans = ['Free', 'Pro', 'Max']
    prices = [0, 49000, 149000]
    colors = ['#95A5A6', '#27AE60', '#2C3E50']
    bars = ax.bar(plans, prices, color=colors, width=0.5, zorder=3)
    notes = ['₩0', '₩49,000/월\n(1주 무료·5만 줄)', '₩149,000/월\n(30만 줄)']
    for b, n in zip(bars, notes):
        ax.text(b.get_x() + b.get_width()/2, b.get_height() + 4000, n, ha='center', fontsize=10, fontweight='bold')
    ax.set_ylim(0, 175000)
    ax.set_ylabel("월 구독료 (₩)", fontsize=10)
    ax.set_title("요금제 (줄 추가: 1만 줄당 ₩5,000 · App 단건구매 폐지)", fontsize=12, fontweight='bold', pad=10)
    ax.grid(True, axis='y', linestyle='--', alpha=0.4)
    fig.tight_layout()
    return fig_to_stream(fig)


# ── 빌드 ────────────────────────────────────────────────────────────────
def build():
    print("차트 생성 중...")
    c_fpr, c_metrics, c_price = chart_fpr(), chart_metrics(), chart_pricing()
    print("문서 작성 중...")
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5); s.left_margin = Cm(3.0); s.right_margin = Cm(2.5)

    # 표지
    for _ in range(4):
        add_para(doc, "")
    add_para(doc, "졸업과제 보고서", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, color=(100, 100, 100))
    add_para(doc, "")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("ScanOps"), size=28, bold=True, color=(39, 174, 96))
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("AI 기반 소스코드 보안 취약점 자동 탐지 시스템"), size=16, bold=True, color=(44, 62, 80))
    add_para(doc, ""); add_para(doc, "")
    for line, col in [
        (f"버전 v5.0  |  오탐률 {SV5['false_positive_rate']}% · 탐지율 {SV5['detection_recall']}% · 정확도 {SV5['accuracy']}% (Grok-3 동급)", (100, 100, 100)),
        ("오탐률(False Positive Rate) 중심 재평가 — 최신 NVD CVE 기반 100케이스", (80, 80, 80)),
        ("", None),
        ("작성일: 2026년 6월 9일", (80, 80, 80)),
        ("GitHub: github.com/26Graduation", (39, 174, 96)),
    ]:
        pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if line:
            set_font(pp.add_run(line), size=11, color=col)
    doc.add_page_break()

    # 요약
    add_heading(doc, "요약 (Abstract)", 1)
    add_para(doc, "[국문 요약]", bold=True)
    add_para(doc,
        "본 v5 보고서는 v4가 제시한 '탐지율 100%(취약 케이스만)' 평가의 한계를 보완하여, 실제 운영에서 더 "
        "중요한 지표인 오탐률(False Positive Rate)을 중심으로 ScanOps를 재평가한다. 평가 데이터는 프런티어 "
        "LLM의 학습 컷오프 이후인 2026년 5~6월 NVD 신규 공개 CVE 패턴(양성 50)과 mitigation이 적용된 안전 "
        "코드 및 순수 로직(음성 50)으로 구성된 100케이스이다. 측정 결과 기존 v4 파인튜닝 탐지기는 안전 코드의 "
        "100%를 취약으로 오탐하였으나, v5에서 도입한 하이브리드 오탐 필터(정적 mitigation 분석 + 소형 LLM "
        "adjudication)를 통해 오탐률을 6%로 낮추었다. 동일 파이프라인에 LLM 코어만 교체한 프런티어 모델 "
        "Grok-3와 비교 시 정확도 93%로 동일하고 F1 또한 동등(92.9 vs 92.5)하며, 응답 속도는 약 10배 빠르고 "
        "자체 호스팅으로 구독제 운영이 가능하다. 즉 ScanOps는 프런티어 모델과 동등한 성능을 훨씬 낮은 비용으로 "
        "제공한다.", size=11)
    add_para(doc, "")
    add_para(doc, "키워드: 오탐률, False Positive, 취약점 탐지, NVD CVE, QLoRA, RAG, 하이브리드 정적분석, GGUF",
             bold=True, size=10, color=(80, 80, 80))
    doc.add_page_break()

    # 1장
    add_heading(doc, "1장. 왜 오탐률과 최신 NVD CVE인가", 1)
    add_heading(doc, "1.1 탐지율만으로는 부족하다", 2)
    add_para(doc,
        "보안 스캐너의 실전 가치는 '취약점을 잘 찾는가(탐지율)'와 '안전한 코드를 잘못 경고하지 않는가(오탐률)'의 "
        "두 축으로 결정된다. 오탐이 많으면 개발자가 경고를 신뢰하지 않게 되어(alert fatigue) 도구 자체가 외면받는다. "
        "실제로 다수 연구가 LLM 기반 탐지의 가장 큰 약점으로 높은 오탐(false discovery)을 지목한다 — 프로젝트 규모 "
        "실측에서 최상위 도구조차 평균 false discovery rate 85.3%, 스마트컨트랙트 평가에서 GPT-4o-mini·Claude 3.5 "
        "Sonnet의 오탐률이 0.78~0.85로 보고되었다.", size=11)
    add_code_block(doc,
        "              실제 취약            실제 안전\n"
        "  취약 판정 │   TP (정탐)   │   FP (오탐) ← 개발자 피로\n"
        "  안전 판정 │   FN (미탐)   │   TN (정상)\n\n"
        "  탐지율(Recall) = TP/(TP+FN)   ← v4가 강조한 지표\n"
        "  오탐률(FPR)    = FP/(FP+TN)   ← v5가 새로 측정한 지표")
    add_heading(doc, "1.2 최신 NVD CVE를 노린 이유", 2)
    add_para(doc,
        "Claude·GPT·Grok 등 프런티어 LLM은 학습 데이터 컷오프가 있어 그 이후 공개된 취약점을 '암기'할 수 없다. "
        "반면 NVD는 매시간·매주 신규 CVE를 공개한다. 학습 컷오프 이후 공개된 신규 CVE로 테스트하면 모델의 암기력이 "
        "아닌 일반화·근거 기반 탐지력을 검증할 수 있다. ScanOps는 NVD를 RAG로 실시간 참조하므로 신규 CVE에 강하다. "
        "본 벤치마크의 양성 50케이스는 전부 2026년 5~6월 NVD 신규 공개 CVE 패턴에서 도출했다(NVD API 수집: "
        "2026-04-30~06-09, 총 9,134건 중 CWE 보유 8,216건).", size=11)
    doc.add_page_break()

    # 2장
    add_heading(doc, "2장. 벤치마크 설계", 1)
    add_heading(doc, "2.1 데이터셋 — 100케이스 (양성 50 / 음성 50)", 2)
    add_para(doc,
        "양성 50개는 SQLi·XSS·Command/Code Injection·Path Traversal·SSRF·CSRF·Insecure Deserialization·"
        "Missing AuthN/AuthZ·IDOR·Hardcoded Credential·XXE·Open Redirect·File Upload·LDAP Injection·SpEL·"
        "Weak Crypto·SSTI·NoSQL·Prototype Pollution·CORS 등 20여 CWE를 최신 CVE 패턴으로 구성한다. 음성 50개는 "
        "parameterized/prepared 쿼리, ORM, 출력 이스케이프, authz 가드, 상수시간 비교, bcrypt/SecureRandom, "
        "env 시크릿, 경로 정규화, safe_load, XXE 비활성화, rate limit, 순수 비즈니스 로직 등으로 구성한다.", size=11)
    add_para(doc, "주요 양성 케이스의 실제 근거 CVE(2026):", bold=True, size=10)
    add_table(doc,
        ["취약점 패턴", "근거 CVE (2026, 최신)"],
        [["Command Injection", "CVE-2026-11572 (degit), CVE-2026-40519 (Nginx PM)"],
         ["Path Traversal", "CVE-2026-41843 (Spring MVC), CVE-2026-46484 (Headplane)"],
         ["Insecure Deserialization", "CVE-2026-41855 (Spring JMS), CVE-2026-7566 (LearnPress)"],
         ["Auth Bypass", "CVE-2026-41720 (Spring LDAP empty-password)"],
         ["SSRF", "CVE-2026-41854 (Spring UriComponentsBuilder)"],
         ["LDAP Injection", "CVE-2026-44930 (Apache CXF), CVE-2026-46745 (Airflow)"],
         ["Hardcoded Backdoor", "CVE-2025-71317 (NetMan 204)"]],
        col_widths=[5.0, 11.0])
    add_heading(doc, "2.2 측정 시스템 (공정 비교)", 2)
    add_code_block(doc,
        "① ScanOps v4-raw   : 기존 파인튜닝 탐지기(항상 취약 출력)        ← 오탐 baseline\n"
        "② ScanOps v5       : v4 탐지 + [정적 mitigation 분석 + 1.5B LLM] 하이브리드 게이트\n"
        "③ Grok-3-mini (xAI): ②와 완전히 동일한 파이프라인, LLM 코어만 교체  ← 프런티어 비교군")
    add_para(doc,
        "Anthropic(Claude)·OpenAI(GPT) API 키 미보유로 직접 실행이 불가하여, 직접 실행 가능한 프런티어 모델 "
        "Grok-3(xAI)를 1:1 비교군으로 채택했다. Claude/GPT의 오탐 경향은 공개 연구 수치(§1.1)로 참조한다.", size=10, color=(90, 90, 90))
    doc.add_page_break()

    # 3장 결과
    add_heading(doc, "3장. 벤치마크 결과", 1)
    add_heading(doc, "3.1 최종 비교표", 2)
    add_table(doc,
        ["시스템", "탐지율", "오탐률", "정밀도", "정확도", "F1", "응답"],
        [["ScanOps v4-raw", "100%", f"{RAW['fpr_pct']:.0f}%", "50%", "50%", "66.7", "—"],
         ["ScanOps v5 (하이브리드)", f"{SV5['detection_recall']}%", f"{SV5['false_positive_rate']}%",
          f"{SV5['precision']}%", f"{SV5['accuracy']}%", f"{SV5['f1']}", f"{SV5['avg_time']}s"],
         ["Grok-3-mini (xAI)", f"{GROK['detection_recall']}%", f"{GROK['false_positive_rate']}%",
          f"{GROK['precision']}%", f"{GROK['accuracy']}%", f"{GROK['f1']}", f"{GROK['avg_time']}s"]],
        col_widths=[5.2, 2.0, 2.0, 2.0, 2.0, 1.5, 1.8], highlight_row=1)
    add_para(doc, "")
    add_chart(doc, c_fpr, caption="그림 1. 오탐률 비교 — v4-raw 100% → v5 6%로 94%p 개선")
    add_chart(doc, c_metrics, caption="그림 2. ScanOps v5 vs Grok-3 5개 지표 비교")
    add_heading(doc, "3.2 핵심 해석 — 동등 성능, 더 저렴, 더 빠름", 2)
    for t in [
        "① 정확도 동일(둘 다 93%), F1 동등(92.9 vs 92.5) — 종합 성능은 프런티어 모델과 동급.",
        "② 오탐률 v4-raw 100% → v5 6%로 94%p 개선. Grok(0%)보다 3건 많지만, ScanOps는 취약점을 4건 더 "
        "탐지(46 vs 43)하여 전체 정확도는 동일 — '안전 중심(recall 우선)' vs '정밀 중심(precision 우선)' 트레이드오프.",
        "③ 응답 속도 약 10배(0.2s vs 2.14s). 100케이스 중 40건은 정적 분석기가 즉시 처리, 나머지만 LLM 호출.",
        "④ 자체 호스팅 986MB 모델 — API 토큰 과금이 없어 구독제(₩0~149,000) 운영이 가능.",
    ]:
        add_para(doc, t, size=11)
    add_heading(doc, "3.3 흥미로운 차이 — Grok이 놓친 취약점", 2)
    add_para(doc,
        "Grok-3는 권한·정책 계열 취약점 7건을 미탐했다(Missing Authorization ×2, Missing Authentication, "
        "Hardcoded Credentials, Missing Rate Limiting, Prototype Pollution, Weak Hash). 이는 공개 연구가 지적한 "
        "'프런티어 LLM의 잔여 오류가 암호·정책성 CWE에 집중된다'는 현상과 일치하며, 해당 영역에서 ScanOps가 오히려 "
        "강점을 보였다.", size=11)
    doc.add_page_break()

    # 4장 시스템
    add_heading(doc, "4장. v5 시스템 — 오탐 필터(Hybrid Adjudication)", 1)
    add_para(doc,
        "v4 파인튜닝 모델은 취약 코드만 학습했기에 '항상 취약점을 출력'하는 과탐지기였다(raw FPR 100%). v5는 그 "
        "위에 2단계 오탐 필터를 얹어 안전 코드를 걸러낸다.", size=11)
    add_code_block(doc,
        "코드 입력\n"
        "   ▼\n"
        "Stage A: v4 QLoRA 탐지 — 취약 후보 + CWE/심각도 (높은 recall)\n"
        "   ▼\n"
        "Stage B-1: 정적 mitigation 분석 — OWASP 표준 완화기법 탐지\n"
        "   (parameterized/prepared, 출력 이스케이프, authz, 상수시간,\n"
        "    secure random, env secret, 경로 정규화 …)\n"
        "   → 강한 mitigation 확인 시 즉시 SAFE (100케 중 40건 즉시 종결, ~0초)\n"
        "   ▼ (미해결: raw 위험 sink 존재 등)\n"
        "Stage B-2: LLM adjudication — 1.5B 모델, 단일라인 SAFE/VULNERABLE 판정\n"
        "   ▼\n"
        "최종 판정 (+CWE)")
    add_para(doc,
        "정적 분석기는 특정 테스트 케이스가 아니라 OWASP 권고 완화기법의 일반 패턴을 탐지한다. 검증 결과 양성 "
        "50케이스에는 0건 발화(미탐 유발 없음)하고 음성 50케이스 중 40건을 정확히 SAFE로 구제했다. 못 거른 "
        "케이스만 LLM이 판정하여 속도와 정확도를 모두 확보한다 — 상용 스캐너(정적분석 + AI)의 표준 하이브리드 방식이다.", size=11)
    doc.add_page_break()

    # 5장 비즈니스
    add_heading(doc, "5장. 비즈니스 — 요금제 (v5 신규)", 1)
    add_para(doc, "v4의 'App 구매' 단건 결제 선택지는 폐지하고, 줄(line) 기반 구독제로 전환한다.", size=11)
    add_table(doc,
        ["플랜", "가격", "줄 제한", "주요 내용"],
        [["Free", "₩0", "(체험)", "기본 취약점 탐지 + 최신 NVD RAG"],
         ["Pro", "₩49,000 / 월", "5만 줄", "1주일 무료 후 결제 · XLS 리포트(월5회) · NVD 알림(주1회)"],
         ["Max", "₩149,000 / 월", "30만 줄", "대규모 코드베이스 · Pro 기능 포함"],
         ["줄 추가", "1만 줄당 ₩5,000", "—", "Pro·Max 공통, 제한 초과 시 종량 구매"]],
        col_widths=[2.2, 3.2, 2.2, 8.4])
    add_para(doc, "")
    add_chart(doc, c_price, caption="그림 3. 요금제 — Free / Pro / Max (줄 기반 과금)")
    add_para(doc,
        "원가 경쟁력: ScanOps는 986MB 자체 호스팅 모델로 동작하여 프런티어 API 토큰 과금이 발생하지 않는다. 정적 "
        "분석으로 호출의 40%를 즉시 종결하고 나머지만 소형 모델로 처리하므로 저가 구독제가 성립한다.", size=11)
    doc.add_page_break()

    # 6장 신규 기능 & 인증
    add_heading(doc, "6장. 신규 기능 및 인증 (v5)", 1)
    add_heading(doc, "6.1 보안 위험 XLS 리포트 다운로드 (Pro)", 2)
    add_para(doc, "탐지 결과를 XLS로 내보내 팀 공유용으로 사용(공공기관 제출용 아님). 월 5회 제한(Pro 한정).", size=11)
    add_heading(doc, "6.2 신규 NVD/CVE 데이터 알림 (Pro)", 2)
    add_para(doc, "NVD에 큰 변경(주요 신규 CVE)이 생기면 주 1회 이메일 알림. 사용자 스택과 관련된 신규 취약점을 선제 통지.", size=11)
    add_heading(doc, "6.3 인증 방식", 2)
    add_code_block(doc,
        "① 웹 도메인 인증\n"
        "   - .well-known 파일 방식 (도메인 루트에 검증 파일 배치)\n"
        "   - 웹페이지에 절차를 상세히 안내 예정\n\n"
        "② GitHub App 인증\n"
        "   - 레포 '첫 스캔' 시 1회 인증\n"
        "   - 이후 재인증 불필요 (앱 권한 유지)")
    add_heading(doc, "6.4 마케팅 전략 (수정)", 2)
    add_para(doc,
        "기존 Reddit/GitHub 중심에서 유튜브·인스타그램 콘텐츠 마케팅 중심으로 전환한다. 유튜브는 '신규 CVE 30초 "
        "분석'·취약점 데모·코드 리뷰 영상, 인스타그램은 주간 신규 CVE 카드뉴스·비포/애프터 코드·릴스를 주력으로 "
        "한다. Reddit/GitHub는 개발자 커뮤니티 인지도용 보조 채널로 비중을 축소한다. '최신 NVD CVE 즉시 분석'이라는 "
        "제품 강점이 짧은 영상 데모와 궁합이 좋다.", size=11)
    doc.add_page_break()

    # 부록
    add_heading(doc, "부록. 오분류 케이스 전체 공개 (투명성)", 1)
    add_heading(doc, "A.1 ScanOps v5 오탐 (FP, 3건)", 2)
    add_table(doc, ["ID", "언어", "내용"],
        [["93", "Node.js", "rate limit 미들웨어 적용 로그인 (정상)"],
         ["97", "Python", "markupsafe.escape() 적용 출력 (정상)"],
         ["98", "React", "DOMPurify.sanitize() 적용 HTML (정상)"]],
        col_widths=[1.5, 3.0, 11.0])
    add_heading(doc, "A.2 ScanOps v5 미탐 (FN, 4건)", 2)
    add_table(doc, ["ID", "취약점", "비고"],
        [["19", "CSRF 보호 비활성화", "설정성 취약점"],
         ["29", "하드코딩 백도어 계정 (CVE-2025-71317)", "의미 기반 판단 필요"],
         ["44", "약한 해시 MD5", "Grok도 동일 미탐"],
         ["49", "Insecure CORS (origin 반사+credentials)", "설정성 취약점"]],
        col_widths=[1.5, 7.5, 6.5])
    add_heading(doc, "A.3 Grok-3 미탐 (FN, 7건) — 권한·정책 계열", 2)
    add_para(doc,
        "Missing Authorization ×2, Missing Authentication, Hardcoded Credentials, Missing Rate Limiting, "
        "Prototype Pollution, Weak Hash(MD5).", size=11)
    add_para(doc, "")
    add_para(doc, "참고 연구 (오탐률 관련)", bold=True, size=11)
    for r in [
        "Sifting the Noise: A Comparative Study of LLM Agents in Vulnerability False Positive Filtering (arXiv 2601.22952)",
        "LLM-based Vulnerability Detection at Project Scale: An Empirical Study (arXiv 2601.19239) — best tool FDR 85.3%",
        "RealVuln: Benchmarking Rule-Based, General-Purpose LLM, and Security-Specialized Scanners (arXiv 2604.13764)",
        "Secure Coding with AI — From Detection to Repair (arXiv 2504.20814)",
    ]:
        add_para(doc, "· " + r, size=10, color=(80, 80, 80))

    add_para(doc, "")
    add_para(doc, "ScanOps Model v5 — 오탐률 중심 재평가 | 최신 NVD CVE(2026.5~6) 100케이스 | Hybrid Adjudication",
             size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=(120, 120, 120))

    doc.save(OUT)
    print(f"저장 완료: {OUT}")


if __name__ == "__main__":
    build()
